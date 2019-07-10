# _*_ coding:utf-8 _*_

"""
Author:sea85
Data:2019/7/2
将srt字幕文件转换为PDF
"""

import os
import re
from docx import Document
from docx.shared import Inches,Pt,RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

def get_name(str):
    name = str
    res1 = re.search(r'(.*)?(S[\d]{1,2}E[\d]{1,2})',str,flags=re.IGNORECASE)
    res2 = re.search(r'(.*)?(\.[\d]{4}\.[\d]{3,4}p)',str,flags=re.IGNORECASE)
    if res1:
        name = res1.group(1)+res1.group(2)
        print(res1.group(1)+res1.group(2))
        return name
    if res2:
        name = res2.group(1)
        print(res2.group(1))
        return name
    #if res.group(2):

    #res = re.search(r'(.*)?([\d]{3,4}p)',str,flags=re.IGNORECASE)

def get_file_path(path):
    """
    遍历指定目录，返回完整目录列表。
    """
    fullpaths = []
    for dirpath, dirnames, filenames in os.walk(path):
        fullpaths = [os.path.join(dirpath,file) for file in filenames]
    return fullpaths

def isdialog(str,row=0):
    str = str.replace("{\\r译文字幕}","")
    res = re.search(r'Dialogue: 0.*,0,0,0,,(.*)\\N{.*}(.*)',str)
    if res:
        if not re.search(r'[\{\}]',res.group(1)):   #group1无{或}号方为匹配
            #print(res.groups())
            print(res.group(2))
            print(res.group(1))
            style1 = document.styles['Normal']
            font1 = style1.font
            font1.name = "Cambria"
            font1.size = Pt(16)
            style2 = document.styles['Body Text']
            font2 = style2.font
            font2.name = u"微软雅黑"
            style2._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
            font2.size = Pt(14)
            font2.color.rgb = RGBColor(89,89,89)
            p1 = document.add_paragraph(res.group(2),style=style1)
            p1_format = p1.paragraph_format
            #p1_format.space_before = Pt(14)
            p1_format.space_after = Pt(2)
            p2 = document.add_paragraph(res.group(1),style=style2)
            p2_format = p2.paragraph_format
            p2_format.space_before = Pt(2)
            p2_format.space_after = Pt(13)
            #document.add_paragraph('\n')


path = "C:\\Users\Administrator\\Desktop\Personal\\字幕文件"
filepath = "C:\\Users\\Administrator\\Desktop\Personal\\字幕文件\\Chernobyl.S01E01.1.23.45.720p.AMZN.WEB-DL.DDP5.1.H.264-NTb.简体&英文.ass"

document = Document()
#style = document.styles['Normal']
#font = style.font
#font.name = "Cambria"
#font.size = Pt(18)


lists = os.listdir(path)
for list in lists:
    name = get_name(list)
    filepath = os.path.join(path,list)
    #print(name)
    document = Document()
    heading = document.add_heading(name, 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph('  ')

    try:
        with open(filepath,'r',encoding='utf-16',errors='ignore') as f:
            for line in f:
                isdialog(line)
    except UnicodeError as e:
        print(e)
        with open(filepath,'r',encoding='utf-8',errors='ignore') as f:
            for line in f:
                isdialog(line)
    docx_name = name + ".docx"
    document.save(docx_name)
