# _*_ coding:utf-8 _*_

"""
Author:sea85
Data:2019/7/2
将srt字幕文件转换为PDF
"""

import os
import re
import shutil
from docx import Document
from docx.shared import Inches,Pt,RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

path = os.path.join(os.getcwd(),"字幕文件")

def get_name(str):
    name_dict = {}
    name = str
    res1 = re.search(r'(.*)?(S[\d]{1,2}E[\d]{1,2})',str,flags=re.IGNORECASE)
    res2 = re.search(r'(.*)?(\.[\d]{4}\.[\d]{3,4}p)',str,flags=re.IGNORECASE)
    res3 = re.search(r'(.*)?(\.[\d]{3,4}p)',str,flags=re.IGNORECASE)
    if res1:
        name = res1.group(1)+res1.group(2)
        #print(res1.group(1)+res1.group(2))
        name_dict[name] = str
        return name_dict
    if res2:
        name = res2.group(1)
        #print(res2.group(1))
        name_dict[name] = str
        return name_dict
    if res3:
        name = res3.group(1)
        #print(res3.group(1))
        name_dict[name] = str
        return name_dict
    #if res.group(2):

    #res = re.search(r'(.*)?([\d]{3,4}p)',str,flags=re.IGNORECASE)

def move_file(current_path,target_path):
    filepath,file_name = os.path.split(current_path)
    if os.path.exists(os.path.join(target_path,file_name)):
        os.remove(os.path.join(target_path,file_name))
    shutil.move(current_path,target_path)



def isdialog(str,document,sub_document):
    str = str.replace("{\\r译文字幕}","")
    #str = str.replace(r"{\fnSIMHEI\fs22\1c&HFFFFFF&\3c&HFF8000&}", "") #临时
    str = str.replace("{\\r}","") #临时
    #print(str)
    #res = re.search(r'Dialogue: 0.*,0{1,4},0{1,4},0{1,4},,(.*)\\N(.*)',str)
    res = re.search(r'Dialogue: 0.*,0{1,4},0{1,4},0{1,4},,(.*)\\N{.*}(.*)',str)
    if res:
        if not re.search(r'[\{\}]',res.group(1)):   #group1无{或}号方为匹配
            #print(res.groups())
            #print(res.group(2))
            #print(res.group(1))
            style1 = document.styles['Normal']
            font1 = style1.font
            font1.name = "Cambria"
            font1.size = Pt(16)
            style2 = document.styles['Body Text']
            font2 = style2.font
            font2.name = u"微软雅黑"
            style2._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
            font2.size = Pt(14)
            font2.color.rgb = RGBColor(89,89,89)
            p1 = document.add_paragraph(res.group(2),style=style1)
            p11 = sub_document.add_paragraph(res.group(2),style=style1)
            p1_format = p1.paragraph_format
            #p1_format.space_before = Pt(14)
            p1_format.space_after = Pt(2)
            p2 = document.add_paragraph(res.group(1),style=style2)
            p22 = sub_document.add_paragraph(res.group(1),style=style2)
            p2_format = p2.paragraph_format
            p2_format.space_before = Pt(2)
            p2_format.space_after = Pt(13)
            #document.add_paragraph('\n')

def get_script_list():    #判断是否为剧集或单一，剧集返回文件列表
    lists = os.listdir(path)
    name_dict = {}
    for list in lists:
        if ".ass" in list:
            file_dict = get_name(list)
            #print(file_dict)
            name_dict.update(file_dict)
    #print(name_dict)
    return is_same_section(name_dict)

def is_same_section(dict):
    dict1 = {}   #中转字典
    res_dict = {}  #待返回字典
    for key,value in dict.items():   #去除文件名的最后两位，根据名字判断
        if key[0:-2] in dict1.keys():
            dict1[key[0:-2]].append(value)
        else:
            dict1[key[0:-2]] = value.split("%^")
            #print("dict1[key[0:-2]]",dict1)
    #print(dict1)
    for key,value in dict1.items():
        #print(len(value))
        if len(value) > 1:
            key = re.sub(r'S\d{1,2}E',replace_header,key)
            res_dict[key] = value
        else:
            for sub_key, sub_value in dict.items():
                #print(value,sub_value)
                if value[0] == sub_value:
                    res_dict[sub_key] = sub_value.split("%^")
                    #print("进入sub,",sub_value)
                    break
    #print(res_dict)
    return res_dict

def replace_header(matched):
    sec = str(matched.group())
    if sec == "S01E" or sec == "S1E":
        sec = "第一季"
    if sec == "S02E" or sec == "S2E":
        sec = "第二季"
    if sec == "S03E" or sec == "S3E":
        sec = "第三季"
    if sec == "S04E" or sec == "S4E":
        sec = "第四季"
    if sec == "S05E" or sec == "S5E":
        sec = "第五季"
    if sec == "S06E" or sec == "S6E":
        sec = "第6季"
    if sec == "S07E" or sec == "S7E":
        sec = "第七季"
    if sec == "S08E" or sec == "S8E":
        sec = "第八季"
    if sec == "S09E" or sec == "S9E":
        sec = "第九季"
    if sec == "S10E" :
        sec = "第十季"
    if sec == "S11E" :
        sec = "第十一季"
    if sec == "S12E" :
        sec = "第十二季"
    if sec == "S13E" :
        sec = "第十三季"
    if sec == "S14E" :
        sec = "第十四季"
    if sec == "S15E" :
        sec = "第十五季"
    return sec


def main():

    file_dict = get_script_list()
    for file_name,file_path in file_dict.items():
        document = Document()
        #print("file_path:",file_path)
        for file in file_path:   #file_path为文件名列表
            sub_document = Document()
            #print("file:",file)
            header = tuple(get_name(file).keys())[0]
            #print(header)
            full_path = os.path.join(path,file)
            heading = document.add_heading(header, 0) #添加标题
            sub_document.add_heading(header, 0) #添加标题
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            document.add_paragraph('  ')
            print("正在打开%s..." % full_path)
            print("正在转换%s,请等待..." % file)
            try:
                with open(full_path, 'r', encoding='utf-16', errors='ignore') as f:
                    for line in f:
                        isdialog(line,document,sub_document)
            except UnicodeError as e:
                #print(e)
                with open(full_path, 'r', encoding='utf-8', errors='ignore') as f:
                    for line in f:
                        isdialog(line,document,sub_document)
            sub_name = header + ".doc"
            sub_document.save(sub_name)
            sub_current_path = os.path.join(os.getcwd(), sub_name)
            sub_target_path = os.path.join(os.getcwd(), "doc/子文件")
            move_file(sub_current_path, sub_target_path)
            document.add_page_break()
            print("子文件%s转换成功！" % sub_name)
        docx_name = file_name + " 精选热门美剧电影台词 学英语单词口语 中英对照" + ".doc"
        document.save(docx_name)
        current_path = os.path.join(os.getcwd(),docx_name)
        target_path = os.path.join(os.getcwd(),"doc")
        move_file(current_path,target_path)
        print("%s转换成功！"%docx_name)








if __name__ == "__main__":

    main()



